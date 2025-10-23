"""
主路由处理模块
定义应用的所有路由和视图函数
"""
import json
import io
from datetime import datetime
from flask import Blueprint, render_template, request, jsonify, Response, current_app, send_file
from werkzeug.exceptions import RequestEntityTooLarge
from docx import Document
from docx.shared import Inches

from config import get_config
from services import FileService, DifyClient, ValidationService
from utils import get_logger
from utils.exceptions import (
    FileValidationError, FileUploadError, DifyAPIError, 
    ValidationError, ServiceError
)

# 创建蓝图
main_bp = Blueprint('main', __name__)
logger = get_logger(__name__)


@main_bp.route('/')
def index():
    """主页"""
    try:
        logger.info("访问主页")
        return render_template('index.html')
    except Exception as e:
        logger.error(f"主页渲染失败: {str(e)}")
        return jsonify({'error': '页面加载失败'}), 500


@main_bp.route('/upload', methods=['POST'])
def upload():
    """文件上传处理"""
    try:
        logger.info("开始处理文件上传请求")
        
        # 获取请求参数
        file = request.files.get('file')
        review_type = request.form.get('review_type')
        contracting_party = request.form.get('contracting_party')
        category = request.form.get('category')
        
        # 添加详细的调试日志
        logger.debug(f"请求参数 - file: {file}, review_type: {review_type}, contracting_party: {contracting_party}, category: {category}")
        logger.debug(f"request.files: {list(request.files.keys())}")
        logger.debug(f"request.form: {dict(request.form)}")
        
        # 验证请求参数
        ValidationService.validate_upload_request(review_type, contracting_party)
        
        logger.info(f"上传文件: {file.filename}, 类型: {review_type}, 当事方: {contracting_party}, 类别: {category}")
        
        # 验证文件
        config = get_config()
        is_valid, error_msg = ValidationService.validate_file(file, config.MAX_CONTENT_LENGTH)
        if not is_valid:
            logger.warning(f"文件验证失败: {error_msg}")
            return jsonify({'error': error_msg}), 400
        
        # 获取文件信息
        file_info = FileService.get_file_info(file)
        logger.debug(f"文件信息: {file_info}")
        
        # 验证文件类型支持
        if not FileService.validate_file_type_support(file_info['type']):
            raise FileValidationError(f"不支持的文件类型: {file_info['type']}")
        
        # 使用固定的用户ID，与原始程序保持一致
        user_id = 'demo_user'
        
        # 获取Dify客户端
        dify_client = DifyClient(config.DIFY_API_TOKEN, config.DIFY_API_BASE_URL)
        
        try:
            # 上传文件到Dify
            file_id = dify_client.upload_file(file, user_id)
            
            # 创建Dify文件输入格式
            file_input = FileService.create_dify_file_input(file, file_id)
            
            # 准备聊天输入参数
            inputs = {
                'file': file_input,
                'Category': '政策' if review_type == '文件审查' else '合同'
            }
            
            # 如果是合同审查，添加合同方信息
            if review_type == '合同审查' and contracting_party:
                inputs['Contracting_party'] = contracting_party
            
            # 过滤None值
            inputs = FileService.filter_none_values(inputs)
            
            logger.info(f"文件上传成功，准备开始审查: {file.filename}")
            
            # 返回格式与原始程序保持一致
            from flask import Response
            import json
            return Response(json.dumps({
                'inputs': inputs,
                'user_id': user_id
            }, ensure_ascii=False), mimetype='application/json')
            
        finally:
            # 关闭Dify客户端连接
            dify_client.close()
            
    except FileValidationError as e:
        logger.warning(f"文件验证失败: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except FileUploadError as e:
        logger.error(f"文件上传失败: {str(e)}")
        return jsonify({'error': f'文件上传失败: {str(e)}'}), 500
    except ValidationError as e:
        logger.warning(f"参数验证失败: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"上传处理未知错误: {str(e)}")
        return jsonify({'error': '文件上传失败，请稍后重试'}), 500


@main_bp.route('/stream_chat', methods=['POST'])
def stream_chat():
    """流式聊天处理"""
    try:
        logger.info("开始处理流式聊天请求")
        
        # 获取请求数据
        data = request.get_json()
        inputs = data.get('inputs', {})
        user_id = data.get('user_id', '')
        
        # 验证请求参数
        ValidationService.validate_stream_chat_request(inputs, user_id)
        
        logger.info(f"流式聊天输入: {json.dumps(inputs, ensure_ascii=False)}")
        
        # 获取Dify客户端
        config = get_config()
        dify_client = DifyClient(config.DIFY_API_TOKEN, config.DIFY_API_BASE_URL)
        
        def generate_response():
            """生成流式响应"""
            try:
                for chunk in dify_client.chat_stream(inputs, user_id):
                    if chunk.strip():
                        # 直接传递原始流式数据，不进行额外的Unicode解码
                        yield chunk
                        
            except DifyAPIError as e:
                logger.error(f"Dify API错误: {str(e)}")
                error_response = f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"
                yield error_response
            except Exception as e:
                logger.error(f"流式聊天未知错误: {str(e)}")
                error_response = f"data: {json.dumps({'error': '聊天服务暂时不可用'}, ensure_ascii=False)}\n\n"
                yield error_response
            finally:
                # 关闭Dify客户端连接
                dify_client.close()
        
        from flask import stream_with_context
        return Response(
            stream_with_context(generate_response()),
            mimetype='text/event-stream'
               )
        
    except ValidationError as e:
        logger.warning(f"流式聊天参数验证失败: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"流式聊天处理未知错误: {str(e)}")
        return jsonify({'error': '聊天服务暂时不可用，请稍后重试'}), 500


@main_bp.route('/download_word', methods=['POST'])
def download_word():
    """下载Word文档"""
    try:
        logger.info("开始处理Word文档下载请求")
        
        # 获取请求数据
        data = request.get_json()
        if not data:
            return jsonify({'error': '请求数据为空'}), 400
        
        result_type = data.get('result_type')
        result_data = data.get('result_data')
        filename = data.get('filename', 'audit_result.docx')
        
        if not result_type or not result_data:
            return jsonify({'error': '缺少必要参数'}), 400
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = '初审结果' if result_type == 'primary' else '复审结果'
        doc.add_heading(title, 0)
        
        # 添加生成时间
        doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')  # 空行
        
        # 添加审查结果内容
        doc.add_heading('审查结果详情', level=1)
        
        # 如果result_data是字符串，直接添加
        if isinstance(result_data, str):
            doc.add_paragraph(result_data)
        # 如果是字典，格式化显示
        elif isinstance(result_data, dict):
            for key, value in result_data.items():
                p = doc.add_paragraph()
                p.add_run(f'{key}: ').bold = True
                p.add_run(str(value))
        # 如果是列表，逐项显示
        elif isinstance(result_data, list):
            for item in result_data:
                doc.add_paragraph(f'• {str(item)}')
        else:
            doc.add_paragraph(str(result_data))
        
        # 添加页脚信息
        doc.add_paragraph('')
        doc.add_paragraph('---')
        doc.add_paragraph('本文档由文件合同审查系统自动生成')
        
        # 将文档保存到内存
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        logger.info(f"Word文档生成成功: {filename}")
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Word文档生成失败: {str(e)}")
        return jsonify({'error': 'Word文档生成失败'}), 500


@main_bp.route('/health', methods=['GET'])
def health_check():
    """健康检查端点"""
    try:
        logger.debug("执行健康检查")
        
        # 检查配置
        config = get_config()
        
        # 检查Dify API连接
        dify_client = DifyClient(config.DIFY_API_TOKEN, config.DIFY_API_BASE_URL)
        dify_healthy = dify_client.health_check()
        dify_client.close()
        
        health_status = {
            'status': 'healthy' if dify_healthy else 'degraded',
            'services': {
                'dify_api': 'healthy' if dify_healthy else 'unhealthy',
                'file_service': 'healthy',
                'validation_service': 'healthy'
            },
            'timestamp': logger.handlers[0].formatter.formatTime(logger.makeRecord(
                '', 0, '', 0, '', (), None
            )) if logger.handlers else None
        }
        
        status_code = 200 if dify_healthy else 503
        
        logger.info(f"健康检查完成: {health_status['status']}")
        return jsonify(health_status), status_code
        
    except Exception as e:
        logger.error(f"健康检查失败: {str(e)}")
        return jsonify({
            'status': 'unhealthy',
            'error': str(e)
        }), 503


# 错误处理器
@main_bp.errorhandler(413)
def file_too_large(error):
    """文件过大错误处理"""
    logger.warning("文件大小超出限制")
    return jsonify({
        'error': '文件大小超出限制',
        'max_size': f"{current_app.config.get('MAX_CONTENT_LENGTH', 0) // (1024*1024)}MB"
    }), 413


@main_bp.errorhandler(400)
def bad_request(error):
    """错误请求处理"""
    logger.warning(f"错误请求: {str(error)}")
    return jsonify({
        'error': '请求格式错误',
        'message': str(error.description) if hasattr(error, 'description') else '请求无效'
    }), 400


@main_bp.errorhandler(500)
def internal_error(error):
    """内部服务器错误处理"""
    logger.error(f"内部服务器错误: {str(error)}")
    return jsonify({
        'error': '服务器内部错误',
        'message': '服务暂时不可用，请稍后重试'
    }), 500


@main_bp.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(error):
    """处理文件过大异常"""
    logger.warning("上传文件过大")
    return jsonify({
        'error': '上传文件过大',
        'max_size': f"{current_app.config.get('MAX_CONTENT_LENGTH', 0) // (1024*1024)}MB"
    }), 413


# 请求前后处理
@main_bp.before_request
def before_request():
    """请求前处理"""
    logger.debug(f"收到请求: {request.method} {request.path}")


@main_bp.after_request
def after_request(response):
    """请求后处理"""
    logger.debug(f"响应状态: {response.status_code}")
    return response