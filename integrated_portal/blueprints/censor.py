# -*- coding: utf-8 -*-
"""
文件审查系统蓝图
集成文件审查功能到统一门户
"""
import json
import io
import os
import requests
from datetime import datetime
from flask import Blueprint, render_template, render_template_string, request, jsonify, Response, stream_with_context, send_file
from werkzeug.exceptions import RequestEntityTooLarge
from werkzeug.datastructures import FileStorage
from docx import Document
from docx.shared import Inches

# 使用Blueprint整合"文件审查系统"到统一门户
censor_bp = Blueprint(
    'censor', __name__,
    template_folder=os.path.abspath(os.path.join(os.path.dirname(__file__), '../../censor_system/templates')),
    static_folder=os.path.abspath(os.path.join(os.path.dirname(__file__), '../../censor_system/static'))
)

# Dify API配置
DIFY_API_BASE_URL = os.environ.get('DIFY_API_BASE_URL', 'http://localhost/v1')
DIFY_API_TOKEN = os.environ.get('DIFY_API_TOKEN', 'app-h396mRpdDGYuXJ0up68YBqQP')

# 文件类型映射
EXT_TYPE_MAP = {
    'application/pdf': 'pdf',
    'application/msword': 'doc',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'text/plain': 'txt',
    'text/markdown': 'markdown',
    'text/html': 'html',
    'application/vnd.ms-excel': 'xls',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
    'application/vnd.ms-powerpoint': 'ppt',
    'application/vnd.openxmlformats-officedocument.presentationml.presentation': 'pptx',
    'application/xml': 'xml',
    'application/epub+zip': 'epub',
    'image/jpeg': 'jpg',
    'image/png': 'png',
    'image/gif': 'gif',
    'image/webp': 'webp',
    'image/svg+xml': 'svg',
    'audio/mpeg': 'mp3',
    'audio/mp4': 'm4a',
    'audio/wav': 'wav',
    'audio/webm': 'webm',
    'audio/amr': 'amr',
    'video/mp4': 'mp4',
    'video/quicktime': 'mov',
    'video/mpeg': 'mpeg',
    'audio/mpga': 'mpga',
}

EXT_FILETYPE_MAP = {
    'pdf': 'document', 'doc': 'document', 'docx': 'document', 'txt': 'document', 'md': 'document', 'markdown': 'document', 'html': 'document',
    'xls': 'document', 'xlsx': 'document', 'ppt': 'document', 'pptx': 'document', 'xml': 'document', 'epub': 'document',
    'csv': 'document', 'eml': 'document', 'msg': 'document',
    'jpg': 'image', 'jpeg': 'image', 'png': 'image', 'gif': 'image', 'webp': 'image', 'svg': 'image',
    'mp3': 'audio', 'm4a': 'audio', 'wav': 'audio', 'webm': 'audio', 'amr': 'audio', 'mpga': 'audio',
    'mp4': 'video', 'mov': 'video', 'mpeg': 'video',
}

# 配置常量
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB
API_TIMEOUT = 120
SUPPORTED_EXTENSIONS = {
    'pdf', 'doc', 'docx', 'txt', 'md', 'markdown', 'html',
    'xls', 'xlsx', 'ppt', 'pptx', 'xml', 'epub', 'csv'
}

def guess_type(file):
    """推断文件类型"""
    mime = getattr(file, 'mimetype', None)
    ext = EXT_TYPE_MAP.get(mime)
    if not ext:
        ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    return EXT_FILETYPE_MAP.get(ext, 'document')

def create_dify_file_input(file, file_id):
    """创建Dify文件输入格式"""
    file_type = guess_type(file)
    return {
        'type': file_type,
        'upload_file_id': file_id,
        'transfer_method': 'local_file'
    }

def filter_none_values(obj):
    """过滤None值"""
    if isinstance(obj, dict):
        return {k: filter_none_values(v) for k, v in obj.items() if v is not None}
    elif isinstance(obj, list):
        return [filter_none_values(i) for i in obj if i is not None]
    else:
        return obj

def validate_file(file, max_size):
    """验证文件"""
    if not file or not file.filename:
        return False, '请选择文件'
    
    # 检查文件大小
    file.seek(0, 2)  # 移动到文件末尾
    file_size = file.tell()
    file.seek(0)  # 重置到开头
    
    if file_size > max_size:
        return False, f'文件大小超出限制 ({max_size // (1024*1024)}MB)'
    
    # 检查文件扩展名
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in SUPPORTED_EXTENSIONS:
        return False, f'不支持的文件类型: {ext}'
    
    return True, None

def validate_upload_request(review_type, contracting_party):
    """验证上传请求参数"""
    if not review_type:
        raise ValueError('审查类型不能为空')
    
    if review_type not in ['文件审查', '合同审查']:
        raise ValueError('无效的审查类型')
    
    if review_type == '合同审查' and not contracting_party:
        raise ValueError('合同审查需要指定合同方')

@censor_bp.route('/', methods=['GET'])
def index():
    """主页"""
    # 直接读取原子系统模板内容渲染
    tpl_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../censor_system/templates/index.html'))
    with open(tpl_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return render_template_string(content)

@censor_bp.route('/upload', methods=['POST'])
def upload():
    """文件上传处理"""
    try:
        # 获取请求参数
        file = request.files.get('file')
        review_type = request.form.get('review_type')
        contracting_party = request.form.get('contracting_party')
        category = request.form.get('category')
        
        # 验证请求参数
        validate_upload_request(review_type, contracting_party)
        
        # 验证文件
        is_valid, error_msg = validate_file(file, MAX_CONTENT_LENGTH)
        if not is_valid:
            return jsonify({'error': error_msg}), 400
        
        # 使用固定的用户ID
        user_id = 'demo_user'
        
        try:
            # 上传文件到Dify
            url = f"{DIFY_API_BASE_URL}/files/upload"
            headers = {"Authorization": f"Bearer {DIFY_API_TOKEN}"}
            
            # 确保文件指针在开头
            file.seek(0)
            
            files_data = {'file': (file.filename, file.stream, file.mimetype)}
            data = {'user': user_id}
            
            resp = requests.post(url, headers=headers, files=files_data, data=data, timeout=API_TIMEOUT)
            resp.raise_for_status()
            
            file_info = resp.json()
            file_id = file_info['id']
            
            # 创建Dify文件输入格式
            file_input = create_dify_file_input(file, file_id)
            
            # 准备聊天输入参数
            inputs = {
                'file': file_input,
                'Category': '政策' if review_type == '文件审查' else '合同'
            }
            
            # 如果是合同审查，添加合同方信息
            if review_type == '合同审查' and contracting_party:
                inputs['Contracting_party'] = contracting_party
            
            # 过滤None值
            inputs = filter_none_values(inputs)
            
            # 返回格式与原始程序保持一致
            return Response(json.dumps({
                'inputs': inputs,
                'user_id': user_id
            }, ensure_ascii=False), mimetype='application/json')
            
        except requests.exceptions.RequestException as e:
            return jsonify({'error': f'文件上传失败: {str(e)}'}), 500
            
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': '文件上传失败，请稍后重试'}), 500

@censor_bp.route('/stream_chat', methods=['POST'])
def stream_chat():
    """流式聊天处理"""
    try:
        # 获取请求数据
        data = request.get_json()
        inputs = data.get('inputs', {})
        user_id = data.get('user_id', '')
        
        if not inputs or not user_id:
            return jsonify({'error': '参数缺失'}), 400
        
        def generate_response():
            """生成流式响应"""
            try:
                url = f"{DIFY_API_BASE_URL}/chat-messages"
                headers = {
                    "Authorization": f"Bearer {DIFY_API_TOKEN}",
                    "Content-Type": "application/json"
                }
                
                payload = {
                    "query": "请审查",
                    "inputs": inputs,
                    "user": user_id,
                    "response_mode": "streaming",
                    "conversation_id": ""
                }
                
                with requests.post(url, headers=headers, json=payload, stream=True, timeout=API_TIMEOUT) as resp:
                    for line in resp.iter_lines():
                        if line:
                            if line.startswith(b'data:'):
                                content = line[5:].decode('utf-8').strip()
                                yield f"data: {content}\n\n"
                                
            except requests.exceptions.RequestException as e:
                error_response = f"data: {json.dumps({'error': f'Dify API错误: {str(e)}'}, ensure_ascii=False)}\n\n"
                yield error_response
            except Exception as e:
                error_response = f"data: {json.dumps({'error': '聊天服务暂时不可用'}, ensure_ascii=False)}\n\n"
                yield error_response
        
        return Response(
            stream_with_context(generate_response()),
            mimetype='text/event-stream'
        )
        
    except Exception as e:
        return jsonify({'error': '聊天服务暂时不可用，请稍后重试'}), 500

@censor_bp.route('/download_word', methods=['POST'])
def download_word():
    """下载Word文档"""
    try:
        # 获取请求数据
        data = request.get_json()
        if not data:
            return jsonify({'error': '请求数据为空'}), 400
        
        result_type = data.get('result_type')
        result_data = data.get('result_data')
        filename = data.get('filename', 'audit_result.docx')
        
        if not result_type or not result_data:
            return jsonify({'error': '缺少必要参数'}), 400
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = '初审结果' if result_type == 'primary' else '复审结果'
        doc.add_heading(title, 0)
        
        # 添加生成时间
        doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')  # 空行
        
        # 添加审查结果内容
        doc.add_heading('审查结果详情', level=1)
        
        # 如果result_data是字符串，直接添加
        if isinstance(result_data, str):
            doc.add_paragraph(result_data)
        # 如果是字典，格式化显示
        elif isinstance(result_data, dict):
            for key, value in result_data.items():
                p = doc.add_paragraph()
                p.add_run(f'{key}: ').bold = True
                p.add_run(str(value))
        # 如果是列表，逐项显示
        elif isinstance(result_data, list):
            for item in result_data:
                doc.add_paragraph(f'• {str(item)}')
        else:
            doc.add_paragraph(str(result_data))
        
        # 添加页脚信息
        doc.add_paragraph('')
        doc.add_paragraph('---')
        doc.add_paragraph('本文档由文件合同审查系统自动生成')
        
        # 将文档保存到内存
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': 'Word文档生成失败'}), 500

@censor_bp.route('/health', methods=['GET'])
def health_check():
    """健康检查端点"""
    try:
        # 检查Dify API连接
        url = f"{DIFY_API_BASE_URL}/parameters"
        headers = {"Authorization": f"Bearer {DIFY_API_TOKEN}"}
        
        try:
            resp = requests.get(url, headers=headers, timeout=10)
            dify_healthy = resp.status_code == 200
        except:
            dify_healthy = False
        
        health_status = {
            'status': 'healthy' if dify_healthy else 'degraded',
            'services': {
                'dify_api': 'healthy' if dify_healthy else 'unhealthy',
                'file_service': 'healthy',
                'validation_service': 'healthy'
            },
            'timestamp': datetime.now().isoformat()
        }
        
        status_code = 200 if dify_healthy else 503
        return jsonify(health_status), status_code
        
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e)
        }), 503

# 错误处理器
@censor_bp.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(error):
    """处理文件过大异常"""
    return jsonify({
        'error': '上传文件过大',
        'max_size': f"{MAX_CONTENT_LENGTH // (1024*1024)}MB"
    }), 413

# 请求前后处理
@censor_bp.before_request
def before_request():
    """请求前处理"""
    pass

@censor_bp.after_request
def after_request(response):
    """请求后处理"""
    return response

# 错误处理器
@censor_bp.errorhandler(413)
def file_too_large(error):
    """文件过大错误处理"""
    return jsonify({
        'error': '文件大小超出限制',
        'max_size': f"{MAX_CONTENT_LENGTH // (1024*1024)}MB"
    }), 413

@censor_bp.errorhandler(400)
def bad_request(error):
    """错误请求处理"""
    return jsonify({
        'error': '请求格式错误',
        'message': str(error.description) if hasattr(error, 'description') else '请求无效'
    }), 400

@censor_bp.errorhandler(500)
def internal_error(error):
    """内部服务器错误处理"""
    return jsonify({
        'error': '服务器内部错误',
        'message': '服务暂时不可用，请稍后重试'
    }), 500